"""Python port of _render.R — render the thesis via pandoc directly.

Merges index.Rmd + Rmd/*.Rmd (sorted) into one source and renders to docx
using the SJTU Word template, then post-processes for tables, captions,
image wrapping, and per-chapter sections.

Usage:
    python render.py                       # docx (zh, default), CSL=numeric
    python render.py docx
    python render.py docx-en
    python render.py --csl=note            # numeric (default) | note | author-date
    python render.py docx author-date      # bare CSL name also accepted
"""
import copy
import datetime
import glob
import os
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml


# XML namespace URIs used throughout post-processing.
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"


# ── Shared XML helpers  ───────────────────────────────────

def _get_style_maps(doc):
    """Return (name_to_id, heading1_ids) built from doc.styles."""
    name_to_id = {s.name: s.style_id for s in doc.styles}
    heading1_ids = {
        name_to_id.get("Heading 1", "Heading1"),
        "Heading1", "Heading 1", "1",
        name_to_id.get("非编号章节标题"),
        name_to_id.get("非编号章节标题（目录不显示）"),
    }
    heading1_ids.discard(None)
    return name_to_id, heading1_ids


def _ensure_pPr(p):
    """Return the <w:pPr> child of p, creating it if absent."""
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p.insert(0, pPr)
    return pPr


def _set_para_style(p, style_id: str) -> None:
    """Set paragraph style, replacing any existing pStyle."""
    pPr = _ensure_pPr(p)
    for old in pPr.findall(qn("w:pStyle")):
        pPr.remove(old)
    ps = OxmlElement("w:pStyle")
    ps.set(qn("w:val"), style_id)
    pPr.insert(0, ps)


def _get_para_style(p):
    """Return the paragraph style ID, or None."""
    ps = p.find(qn("w:pPr") + "/" + qn("w:pStyle"))
    return ps.get(qn("w:val")) if ps is not None else None


def _get_para_text(p) -> str:
    """Return concatenated text from all <w:t> elements in p."""
    return "".join((t.text or "") for t in p.findall(".//" + qn("w:t")))


def _add_text_run(parent, text: str, preserve: bool = False) -> None:
    """Append a <w:r><w:t>text</w:t></w:r> to parent."""
    r = OxmlElement("w:r")
    t_el = OxmlElement("w:t")
    if preserve:
        t_el.set(qn("xml:space"), "preserve")
    t_el.text = text
    r.append(t_el)
    parent.append(r)


def _make_page_break_para():
    """Build an empty <w:p> containing only <w:br w:type='page'/>."""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


def _append_seq_field(parent, seq_name: str) -> None:
    """Append begin/instrText/end runs for a SEQ field to *parent*."""
    f1 = OxmlElement("w:r")
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    f1.append(fc1)
    parent.append(f1)

    f2 = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {seq_name} \\* ARABIC "
    f2.append(instr)
    parent.append(f2)

    f3 = OxmlElement("w:r")
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    f3.append(fc2)
    parent.append(f3)


def _make_caption_para(label: str, seq_name: str, text: str,
                       caption_style_id: str):
    """Build a caption paragraph: '<label> { SEQ <seq_name> } <text>'."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), caption_style_id)
    pPr.append(pStyle)
    p.append(pPr)
    _add_text_run(p, f"{label} ", preserve=True)
    _append_seq_field(p, seq_name)
    if text:
        _add_text_run(p, f" {text}", preserve=True)
    return p


def _remap_sectpr_rids(sectpr, tpl_rid_to_target, out_target_to_rid):
    """Remap header/footer rIds in a cloned sectPr; drop unmapped refs."""
    for ref in list(sectpr):
        if ref.tag not in (qn("w:headerReference"), qn("w:footerReference")):
            continue
        tpl_rid = ref.get(f"{{{_R_NS}}}id")
        target = tpl_rid_to_target.get(tpl_rid)
        out_rid = out_target_to_rid.get(target) if target else None
        if out_rid:
            ref.set(f"{{{_R_NS}}}id", out_rid)
        else:
            sectpr.remove(ref)


def find_pandoc() -> str:
    py_dir = Path(sys.executable).parent
    # Walk up to find a conda root (one with an `envs/` subdir) and prefer
    # `<root>/envs/pandoc/Library/bin/pandoc.exe`. Works whether the running
    # interpreter is the base env or a sub-env like `envs/mybase`.
    preferred = []
    for anc in [py_dir, *py_dir.parents]:
        cand = anc / "envs" / "pandoc" / "Library" / "bin" / "pandoc.exe"
        if cand.is_file():
            preferred.append(cand)
            break
    preferred += [Path(p) for p in glob.glob(
        str(py_dir / "envs" / "*" / "Library" / "bin" / "pandoc.exe")
    )]
    for c in preferred:
        if c.is_file():
            return str(c)
    exe = shutil.which("pandoc")
    if exe:
        return exe
    candidates = [
        py_dir / "pandoc.exe",
        py_dir / "Scripts" / "pandoc.exe",
        py_dir / "Library" / "bin" / "pandoc.exe",
    ]
    candidates += [Path(p) for p in glob.glob(
        str(py_dir / "envs" / "*" / "Scripts" / "pandoc.exe")
    )]
    for c in candidates:
        if c.is_file():
            return str(c)
    sys.exit("pandoc not found.")


HERE = Path(__file__).resolve().parent
MD_DIR = HERE / "md"
BIB = HERE /  "401-reference.bib"
OUT_BASE = HERE / "SJTUThesis"

# YAML front-matter prepended to the merged source. index.md is no longer
# read — title/author come from the template §1 (see _inject_template_section1);
# only citeproc-relevant metadata remains here.
FRONT_MATTER = (
    "---\n"
    f"bibliography: [{BIB.name}]\n"
    "nocite: '@*'\n"
    "link-citations: yes\n"
    "---\n"
)


def _tag_top_headings(text: str, style: str) -> str:
    # Apply a docx custom-style + .unnumbered to every level-1 ATX heading.
    def repl(m: "re.Match[str]") -> str:
        title = m.group(1).strip()
        # Don't double-tag if already has an attribute block.
        if title.endswith("}"):
            return m.group(0)
        return f'# {title} {{.unnumbered custom-style="{style}"}}'
    return re.sub(r"^#\s+(.+?)\s*$", repl, text, flags=re.MULTILINE)


def merge_sources(docx_mode: bool = False) -> Path:
    parts = [FRONT_MATTER]
    for f in sorted(MD_DIR.glob("*.md")):
        body = f.read_text(encoding="utf-8")
        if docx_mode:
            stem = f.stem
            prefix = stem.split("-", 1)[0] if "-" in stem else stem
            # First non-zero digit of the prefix → 1-based template section.
            # Per-section heading-1 style is applied in post-processing
            # (pandoc ignores custom-style on headings).
            digit = next((c for c in prefix if c.isdigit() and c != "0"), "1")
            # Insert a md-boundary marker so post-processing knows which
            # template section's header/footer applies to each chapter.
            marker = (
                f'\n\n::: {{custom-style="RmdMarker{digit}"}}\n'
                f"§\n"
                f":::\n\n"
            )
            body = marker + body
        parts.append(body)
    merged = "\n\n".join(parts)
    # Strip knitr code chunks (```{r ...} ... ```), leaving inline text intact.
    merged = re.sub(
        r"^```\{[^}]*\}.*?^```\s*$", "", merged,
        flags=re.MULTILINE | re.DOTALL,
    )
    # Replace inline `r format(Sys.Date(), format='...')` with today's date.
    today = datetime.date.today()
    def _eval_inline_r(m: "re.Match[str]") -> str:
        body = m.group(1)
        fmt_match = re.search(
            r"format\s*\(\s*Sys\.Date\(\)\s*,\s*format\s*=\s*['\"]([^'\"]+)['\"]",
            body,
        )
        if fmt_match:
            # R's strftime format tokens are identical to Python's.
            return today.strftime(fmt_match.group(1))
        if "Sys.Date()" in body:
            return today.isoformat()
        return ""
    merged = re.sub(r"`r\s+([^`]*)`", _eval_inline_r, merged)
    tmp = HERE / "_merged.md"
    tmp.write_text(merged, encoding="utf-8")
    return tmp


CSL_CHOICES = {
    "note": HERE / "china-national-standard-gb-t-7714-2015-note.csl",
    "numeric": HERE / "china-national-standard-gb-t-7714-2015-numeric.csl",
    "author-date": HERE / "china-national-standard-gb-t-7714-2015-author-date.csl",
}

def render_docx(pandoc: str, src: Path, lang: str = "zh", csl: str = "numeric") -> None:
    ref = HERE / "template" / (
        "中文毕业设计模板260501.docx" if lang == "zh"
        else "英文毕业设计模板250928.docx"
    )
    out = OUT_BASE.with_suffix(".docx") if lang == "zh" else OUT_BASE.with_name(OUT_BASE.name + "_en").with_suffix(".docx")
    csl_path = CSL_CHOICES.get(csl)
    if csl_path is None:
        sys.exit(f"Unknown CSL '{csl}'. Choices: {', '.join(CSL_CHOICES)}")
    cmd = [
        pandoc, str(src), "-o", str(out),
        "--top-level-division=chapter",
        "--toc", "--toc-depth=3",
        f"--bibliography={BIB}",
        f"--csl={csl_path}",
        "--citeproc",
        f"--reference-doc={ref}",
        "--standalone",
    ]
    print("Running:", " ".join(cmd))
    subprocess.run(cmd, cwd=str(HERE), check=True)
    if lang == "zh":
        _postprocess_docx(out)


def _postprocess_docx(path: Path) -> None:
    """Apply 三线表 + 表格 styles, and insert 题注-styled captions
    above tables and below figures (with auto-incrementing SEQ fields).
    """

    doc = DocxDocument(str(path))
    name_to_id, heading1_ids = _get_style_maps(doc)
    style_names = set(name_to_id.keys())

    TABLE_STYLE = "三线表"
    CELL_PARA_STYLE = "表格"
    CAPTION_STYLE = "Caption"  # fallback
    for s in doc.styles:
        if s.name in ("Caption", "题注"):
            CAPTION_STYLE = s.style_id
            break

    # 1) Tables: apply 三线表 style and 表格 paragraph style to every cell.
    if TABLE_STYLE in style_names:
        for tbl in doc.tables:
            try:
                tbl.style = doc.styles[TABLE_STYLE]
            except KeyError:
                pass
            if CELL_PARA_STYLE in style_names:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            try:
                                p.style = doc.styles[CELL_PARA_STYLE]
                            except KeyError:
                                pass

    body = doc.element.body

    # 2) Tables: ensure each has a 题注 caption above it.
    #    Pandoc emits "TableCaption"-styled paragraphs (with text) above tables;
    #    restyle those to "Caption" and prepend "表 { SEQ Table } ". For tables
    #    without an existing caption, insert an empty one.
    for tbl in body.findall(qn("w:tbl")):
        prev = tbl.getprevious()
        existing_cap = None
        if prev is not None and prev.tag == qn("w:p"):
            ps = prev.find(qn("w:pPr") + "/" + qn("w:pStyle"))
            if ps is not None and ps.get(qn("w:val")) in {
                "TableCaption", "Table Caption", CAPTION_STYLE
            }:
                existing_cap = prev
        if existing_cap is not None:
            # Pull existing text, replace with "表 SEQ <text>" Caption paragraph.
            text = _get_para_text(existing_cap).strip()
            existing_cap.addprevious(_make_caption_para("表", "Table", text, CAPTION_STYLE))
            body.remove(existing_cap)
        else:
            tbl.addprevious(_make_caption_para("表", "Table", "", CAPTION_STYLE))

    # 3) Figures: pandoc emits an inline image inside a "CaptionedFigure"-styled
    #    paragraph; the alt text lives in <wp:docPr @descr>. Insert a sibling
    #    "Caption" paragraph immediately after each such figure.
    for p in list(body.findall(qn("w:p"))):
        if _get_para_style(p) not in {"CaptionedFigure", "Captioned Figure"}:
            continue
        descr = ""
        docpr = p.find(".//{%s}docPr" % _WP_NS)
        if docpr is not None:
            descr = (docpr.get("descr") or docpr.get("title") or "").strip()
        p.addnext(_make_caption_para("图", "Figure", descr, CAPTION_STYLE))

    # 4) Remove pandoc's leftover "ImageCaption" paragraphs (we replaced them).
    for p in list(body.findall(qn("w:p"))):
        if _get_para_style(p) in {"ImageCaption", "Image Caption"}:
            body.remove(p)

    # 5) Convert all inline images to Top-and-Bottom (上下型环绕) anchored images.
    #    wp:anchor children MUST appear in schema order:
    #    simplePos, positionH, positionV, extent, effectExtent, wrap*, docPr,
    #    cNvGraphicFramePr?, graphic. Word will refuse to open the file (or
    #    silently drop the image) if this order is violated.
    wp = lambda tag: f"{{{_WP_NS}}}{tag}"
    for inline in list(body.findall(f".//{wp('inline')}")):
        # Pull required/optional pieces out of the inline element.
        extent = inline.find(wp("extent"))
        effect_extent = inline.find(wp("effectExtent"))
        doc_pr = inline.find(wp("docPr"))
        cnv_gfp = inline.find(wp("cNvGraphicFramePr"))
        graphic = inline.find(f"{{{_A_NS}}}graphic")
        if extent is None or doc_pr is None or graphic is None:
            continue  # malformed; leave as-is

        anchor = parse_xml(
            f'<wp:anchor xmlns:wp="{_WP_NS}"'
            ' distT="0" distB="0" distL="114300" distR="114300"'
            ' simplePos="0" relativeHeight="251658240" behindDoc="0"'
            ' locked="0" layoutInCell="1" allowOverlap="1">'
            '<wp:simplePos x="0" y="0"/>'
            '<wp:positionH relativeFrom="column">'
            '<wp:align>center</wp:align>'
            '</wp:positionH>'
            '<wp:positionV relativeFrom="paragraph">'
            '<wp:posOffset>0</wp:posOffset>'
            '</wp:positionV>'
            '<wp:wrapTopAndBottom/>'
            '</wp:anchor>'
        )
        # Insert extent/effectExtent before wrapTopAndBottom (index 3).
        wrap_el = anchor.find(wp("wrapTopAndBottom"))
        wrap_idx = list(anchor).index(wrap_el)
        insert_at = wrap_idx
        anchor.insert(insert_at, extent)
        insert_at += 1
        if effect_extent is not None:
            anchor.insert(insert_at, effect_extent)
            insert_at += 1
        # docPr / cNvGraphicFramePr / graphic come AFTER wrapTopAndBottom.
        anchor.append(doc_pr)
        if cnv_gfp is not None:
            anchor.append(cnv_gfp)
        anchor.append(graphic)

        inline.getparent().replace(inline, anchor)

    # 6) Bibliography: restyle paragraphs after "参考文献" to "列表段落（无编号）".
    bib_target_id = name_to_id.get("列表段落（无编号）")
    if bib_target_id:
        body_ps = body.findall(qn("w:p"))
        ref_idx = None
        for i, p in enumerate(body_ps):
            if _get_para_style(p) not in heading1_ids:
                continue
            txt = _get_para_text(p).replace(" ", "").replace("\u3000", "")
            if txt == "参考文献":
                ref_idx = i
                break

        if ref_idx is not None:
            for p in body_ps[ref_idx + 1:]:
                if _get_para_style(p) in heading1_ids:
                    break  # next chapter — stop
                _set_para_style(p, bib_target_id)

    # 7) Figure/table notes: paragraphs starting with "注："/"Note:" → 图表注 style.
    #    iter() covers both body-level and table-cell paragraphs in one pass.
    note_style_id = name_to_id.get("图表注")
    if note_style_id:
        _NOTE_PREFIXES = ("注：", "注:", "Note:", "Note：")
        for p in body.iter(qn("w:p")):
            txt = _get_para_text(p).lstrip()
            if any(txt.startswith(pfx) for pfx in _NOTE_PREFIXES):
                _set_para_style(p, note_style_id)

    # 8) Move pandoc's auto-generated TOC to between the abstract and the
    #    body (before the RmdMarker3 paragraph) and rename it 目录. MUST run
    #    before §1 injection so the SDT survives the deletion in step 9.
    _relocate_toc(doc)

    # 9) Inject template §1 verbatim (cover + English cover + originality &
    #    authorization statements), applying field substitutions from
    #    101-setup.md. Replaces everything pandoc rendered before §2.
    template_doc = HERE / "template" / "中文毕业设计模板260501.docx"
    setup_md = HERE / "101-setup.md"
    if template_doc.is_file() and setup_md.is_file():
        _inject_template_section1(
            doc, template_doc, _parse_setup_table(setup_md)
        )

    # 10) Per-Rmd section breaks: each chapter gets the header/footer of the
    #     template's Nth section. RmdMarkerN markers flag chapter ownership.
    if template_doc.is_file():
        _apply_per_rmd_sections(doc, template_doc)

    # 11) Ensure a page break before every level-1 heading (每个 # 之前分页).
    #    Skip the first heading and any heading whose previous paragraph already
    #    starts a new page (carries a sectPr, or contains <w:br type=page/>).
    _ensure_page_break_before_headings(doc)

    # 12) Equations: number display math (单行公式自动编号).
    #     A paragraph is display math if it contains m:oMath and no other text.
    #     Apply '公式' style, center-tab + math + right-tab + (SEQ Equation).
    for p in body.findall(qn("w:p")):
        if not p.findall(".//" + qn("m:oMath")):
            continue
        if _get_para_text(p).strip():  # has non-math text → inline, skip
            continue

        # Apply '公式' style if it exists in the template.
        if "公式" in name_to_id:
            _set_para_style(p, name_to_id["公式"])

        # Set up tab stops: center at ~8cm, right at ~16cm.
        pPr = _ensure_pPr(p)
        for old_tag in ("w:jc", "w:tabs"):
            for old in pPr.findall(qn(old_tag)):
                pPr.remove(old)
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "left")
        pPr.append(jc)
        tabs = OxmlElement("w:tabs")
        pPr.append(tabs)
        for val, pos in (("center", "4536"), ("right", "9072")):
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), val)
            tab.set(qn("w:pos"), pos)
            tabs.append(tab)

        # Insert leading tab (center math), then trailing tab + "(SEQ)".
        r_tab_start = OxmlElement("w:r")
        r_tab_start.append(OxmlElement("w:tab"))
        p.insert(1, r_tab_start)  # after pPr

        r_tab_end = OxmlElement("w:r")
        r_tab_end.append(OxmlElement("w:tab"))
        p.append(r_tab_end)

        _add_text_run(p, "(", preserve=True)
        _append_seq_field(p, "Equation")
        _add_text_run(p, ")", preserve=True)

    # 13) Keywords: paragraphs starting with "关键词"/"keywords" → 关键词 style.
    kw_style_id = name_to_id.get("关键词")
    if kw_style_id:
        _KW_PREFIXES = ("关键词：", "关键词:", "keywords:", "keywords：")
        for p in body.findall(qn("w:p")):
            txt = _get_para_text(p).lstrip().lower()
            if any(txt.startswith(pfx) for pfx in _KW_PREFIXES):
                _set_para_style(p, kw_style_id)

    # 14) settings.xml: set <w:updateFields w:val="true"/> so Word refreshes
    #     TOC / SEQ / page-number fields on first open (打开时自动更新目录).
    upd = doc.settings.element.find(qn("w:updateFields"))
    if upd is None:
        upd = OxmlElement("w:updateFields")
        doc.settings.element.append(upd)
    upd.set(qn("w:val"), "true")

    doc.save(str(path))
    print(f"Post-processed {path.name}: applied 三线表 / 表格 styles, inserted 题注 captions, converted images to 上下型环绕, per-chapter sections.")


def _parse_setup_table(path: Path) -> "dict[str, str]":
    """Parse 101-setup.md's 3-column markdown table → {old_text: new_text}.
    Cells may span multiple physical lines (newlines are preserved within a
    cell), so callers can split on '\\n' to handle multi-line keys.
    """
    text = path.read_text(encoding="utf-8")
    rows = re.findall(r"\|([^|]+)\|([^|]+)\|([^|]+)\|", text)
    out: dict[str, str] = {}
    saw_header = False
    for c1, c2, c3 in rows:
        c1s = c1.strip()
        if not saw_header and "需填写字段" in c1s:
            saw_header = True
            continue
        if all(s.strip() and set(s.strip()) <= {"-", ":"} for s in (c1, c2, c3)):
            continue  # md table separator
        # Collapse internal whitespace too — template is now single-line per
        # field, so any embedded line wraps in the .md cell are cosmetic.
        old = re.sub(r"\s+", " ", c2).strip()
        new = re.sub(r"\s+", " ", c3).strip()
        if old and new and old != new:
            out[old] = new
    return out


def _inject_template_section1(doc, template_path: Path, replacements) -> None:
    """Copy template §1 (everything up to and including the first paragraph
    that ends a section) into `doc`, replacing all body content before the
    first RmdMarkerN paragraph. Apply text substitutions from `replacements`
    (multi-line keys split on '\\n': new value goes into the first matching
    paragraph; subsequent lines are blanked).
    """
    tpl = DocxDocument(str(template_path))
    tpl_body = tpl.element.body
    tpl_rid_to_target = {
        rid: rel.target_ref for rid, rel in tpl.part.rels.items()
    }
    out_target_to_rid = {
        rel.target_ref: rid for rid, rel in doc.part.rels.items()
    }

    # 1) Collect §1 children (deep-copied) up to and including the paragraph
    #    that carries the first sectPr inside its pPr.
    sec1: list = []
    found_terminator = False
    for child in tpl_body:
        if child.tag == qn("w:sectPr"):
            break  # body-level sectPr — no §1 boundary found inline
        clone = copy.deepcopy(child)
        sec1.append(clone)
        if child.tag == qn("w:p"):
            sectPr = child.find(qn("w:pPr") + "/" + qn("w:sectPr"))
            if sectPr is not None:
                clone_sectPr = clone.find(qn("w:pPr") + "/" + qn("w:sectPr"))
                _remap_sectpr_rids(clone_sectPr, tpl_rid_to_target, out_target_to_rid)
                found_terminator = True
                break
    if not found_terminator:
        return

    # 2) Apply text substitutions at single-paragraph granularity. Template
    #    has been flattened so every field sits on its own paragraph; no
    #    multi-paragraph window matching needed. Replace ALL occurrences
    #    (e.g. "20XX" appears in both "20XX年XX月" and "June, 20XX").
    def set_text(p, value: str) -> None:
        first_set = False
        for r in p.findall(qn("w:r")):
            for t in r.findall(qn("w:t")):
                if not first_set:
                    t.text = value
                    t.set(qn("xml:space"), "preserve")
                    first_set = True
                else:
                    t.text = ""
        if not first_set and value:
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = value
            r.append(t)
            p.append(r)

    all_paragraphs = [p for el in sec1 for p in el.iter(qn("w:p"))]
    for old, new in replacements.items():
        old_s, new_s = old.strip(), new.strip()
        if not old_s:
            continue
        for p in all_paragraphs:
            pt = _get_para_text(p)
            if old_s in pt:
                set_text(p, pt.replace(old_s, new_s))

    # 3) In the output: remove all top-level body children up to (but not
    #    including) the first RmdMarkerN paragraph, then prepend §1.
    body = doc.element.body
    first_marker = None
    for c in list(body):
        if c.tag != qn("w:p"):
            continue
        sval = _get_para_style(c)
        if sval and sval.startswith("RmdMarker"):
            first_marker = c
            break
    if first_marker is None:
        return
    for c in list(body):
        if c is first_marker:
            break
        body.remove(c)
    for el in sec1:
        first_marker.addprevious(el)


def _ensure_page_break_before_headings(doc) -> None:
    """Prepend a page break paragraph before EVERY level-1 heading paragraph
    (Heading 1 / 非编号章节标题 / 非编号章节标题（目录不显示）), unconditionally
    — including the first one and including headings already preceded by a
    section break or page break. Idempotent: skip if the immediately preceding
    paragraph is itself an empty page-break-only paragraph that we inserted.
    """
    body = doc.element.body
    name_to_id, heading_ids = _get_style_maps(doc)
    # Also include ABSTRACT style.
    abstract_id = name_to_id.get("ABSTRACT")
    if abstract_id:
        heading_ids.add(abstract_id)

    def is_pure_page_break_para(p) -> bool:
        # A paragraph we inserted: exactly one <w:r> containing only <w:br type=page/>.
        runs = p.findall(qn("w:r"))
        if len(runs) != 1:
            return False
        children = list(runs[0])
        if len(children) != 1:
            return False
        c = children[0]
        return c.tag == qn("w:br") and c.get(qn("w:type")) == "page"

    def is_section_break_para(p) -> bool:
        # A paragraph carrying a sectPr starts a new section (and a new page).
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            return False
        return pPr.find(qn("w:sectPr")) is not None

    headings = []
    for p in body.findall(qn("w:p")):
        if _get_para_style(p) in heading_ids:
            headings.append(p)

    for p in headings:
        prev = p.getprevious()
        while prev is not None and prev.tag != qn("w:p"):
            prev = prev.getprevious()
        if prev is not None and is_pure_page_break_para(prev):
            continue  # already has our page-break paragraph
        if prev is not None and is_section_break_para(prev):
            continue  # section break already starts a new page (每节首章不加分页符)
        p.addprevious(_make_page_break_para())


def _relocate_toc(doc) -> None:
    """Move pandoc's auto-generated TOC (a Word TOC SDT) to between the
    abstract and the first body chapter. Replace its 'Table of Contents'
    heading with '目录' (styled 非编号章节标题（目录不显示）) and prepend a
    page break so the TOC starts on a new page.
    """
    body = doc.element.body
    sdt = body.find(qn("w:sdt"))
    if sdt is None:
        return

    # Insertion point: the RmdMarker3 paragraph (right before Introduction).
    target = None
    for p in body.findall(qn("w:p")):
        ps = p.find(qn("w:pPr") + "/" + qn("w:pStyle"))
        if ps is not None and ps.get(qn("w:val")) == "RmdMarker3":
            target = p
            break
    if target is None:
        return

    # Restyle/relabel the TOC heading inside the SDT.
    name_to_id, _ = _get_style_maps(doc)
    hidden_id = name_to_id.get("非编号章节标题（目录不显示）")
    sdt_content = sdt.find(qn("w:sdtContent"))
    if sdt_content is not None:
        for p in sdt_content.findall(qn("w:p")):
            ps = p.find(qn("w:pPr") + "/" + qn("w:pStyle"))
            if ps is None or ps.get(qn("w:val")) != "TOCHeading":
                continue
            if hidden_id:
                ps.set(qn("w:val"), hidden_id)
            # Replace the first text run with "目录"; clear the rest.
            runs = p.findall(qn("w:r"))
            for i, r in enumerate(runs):
                if i == 0:
                    for t in r.findall(qn("w:t")):
                        t.text = "目 录"
                        for extra in r.findall(qn("w:t"))[1:]:
                            r.remove(extra)
                        break
                else:
                    p.remove(r)
            break

    page_break_p = _make_page_break_para()

    # Detach and reinsert: page-break → SDT → empty anchor paragraph → marker.
    # The trailing empty <w:p> exists so _apply_per_rmd_sections can land §2's
    # sectPr on a paragraph AFTER the TOC (otherwise it walks back past the
    # SDT — a non-w:p element it skips — and lands on the page-break paragraph
    # BEFORE the TOC, which would push the TOC into §3).
    toc_anchor_p = OxmlElement("w:p")
    body.remove(sdt)
    target.addprevious(page_break_p)
    target.addprevious(sdt)
    target.addprevious(toc_anchor_p)


def _apply_per_rmd_sections(doc, template_path: Path) -> None:
    """Insert section breaks at chapter boundaries with header/footer
    references taken from the corresponding template section.
    """
    tpl = DocxDocument(str(template_path))
    tpl_sectprs = [s._sectPr for s in tpl.sections]
    tpl_rid_to_target = {
        rid: rel.target_ref for rid, rel in tpl.part.rels.items()
    }
    out_target_to_rid = {
        rel.target_ref: rid for rid, rel in doc.part.rels.items()
    }

    if not tpl_sectprs:
        return

    body = doc.element.body
    name_to_id, _ = _get_style_maps(doc)
    HEADING1_ID = name_to_id.get("Heading 1", "Heading1")
    UNNUMBERED_ID = name_to_id.get("非编号章节标题")
    UNNUMBERED_HIDDEN_ID = name_to_id.get("非编号章节标题（目录不显示）")
    ABSTRACT_ID = name_to_id.get("ABSTRACT")
    HEADING2_ID = name_to_id.get("Heading 2", "Heading2")
    HEADING3_ID = name_to_id.get("Heading 3", "Heading3")
    APPENDIX_H2_ID = name_to_id.get("附录标题2")
    APPENDIX_H3_ID = name_to_id.get("附录标题3")
    HEADING_STYLES = {HEADING1_ID, "Heading1", "Heading 1", "1"}
    if UNNUMBERED_ID:
        HEADING_STYLES.add(UNNUMBERED_ID)
    if UNNUMBERED_HIDDEN_ID:
        HEADING_STYLES.add(UNNUMBERED_HIDDEN_ID)
    HEADING2_STYLES = {HEADING2_ID, "Heading2", "Heading 2", "2"}
    HEADING3_STYLES = {HEADING3_ID, "Heading3", "Heading 3", "3"}
    # digit → target heading style ID (None = leave Heading 1).
    # 第3节（正文）保留 标题1；其他节使用 非编号章节标题；第5节使用隐藏版本。
    DIGIT_HEADING_STYLE = {
        2: UNNUMBERED_ID,
        3: None,
        4: UNNUMBERED_ID,
        5: UNNUMBERED_HIDDEN_ID,
    }

    # Walk paragraphs in body order; pair each chapter heading with the most
    # recent RmdMarker digit. Collect markers for later removal.
    chapters = []         # [(chapter_para, digit), ...]
    markers_to_remove = []
    current_digit = None
    marker_pat = re.compile(r"^RmdMarker(\d+)$")

    for p in list(body.findall(qn("w:p"))):
        sval = _get_para_style(p)
        if sval:
            m = marker_pat.match(sval)
            if m:
                current_digit = int(m.group(1))
                markers_to_remove.append(p)
                continue
        if sval in HEADING_STYLES and current_digit is not None:
            # Restyle the heading per the digit→style mapping.
            target = DIGIT_HEADING_STYLE.get(current_digit)
            # ABSTRACT 标题特殊处理：使用模板的 "ABSTRACT" 样式。
            heading_text = _get_para_text(p).strip()
            if heading_text.upper() == "ABSTRACT" and ABSTRACT_ID:
                target = ABSTRACT_ID
            if target and sval != target:
                ps = p.find(qn("w:pPr") + "/" + qn("w:pStyle"))
                if ps is not None:
                    ps.set(qn("w:val"), target)
            chapters.append((p, current_digit))
        elif current_digit == 4 and sval is not None:
            # 第4节（附录等）的 ##/### 改用 附录标题2/附录标题3。
            ps = p.find(qn("w:pPr") + "/" + qn("w:pStyle"))
            if ps is not None:
                if sval in HEADING2_STYLES and APPENDIX_H2_ID:
                    ps.set(qn("w:val"), APPENDIX_H2_ID)
                elif sval in HEADING3_STYLES and APPENDIX_H3_ID:
                    ps.set(qn("w:val"), APPENDIX_H3_ID)

    if not chapters:
        for p in markers_to_remove:
            body.remove(p)
        return

    def build_sectpr_for_digit(d):
        idx = max(0, min(d - 1, len(tpl_sectprs) - 1))
        new = copy.deepcopy(tpl_sectprs[idx])
        _remap_sectpr_rids(new, tpl_rid_to_target, out_target_to_rid)
        return new

    # For each chapter starting at index >= 1, insert the *previous* chapter's
    # sectPr in the pPr of the paragraph immediately before this chapter
    # heading. That ends the previous chapter's section.
    marker_set = set(markers_to_remove)
    for i in range(1, len(chapters)):
        p_chap, current_digit = chapters[i]
        prev_digit = chapters[i - 1][1]
        
        if current_digit == prev_digit:
            # Same section: avoid a section break so we don't restart page numbers,
            # but insert a page break so the chapter starts on a new page.
            # Skip if inside TOC (SDT), since _relocate_toc already prepended a page break.
            if p_chap.getparent().tag != qn("w:sdtContent"):
                p_chap.addprevious(_make_page_break_para())
            continue

        prev = p_chap.getprevious()
        while prev is not None and (
            prev.tag != qn("w:p") or prev in marker_set
        ):
            prev = prev.getprevious()
        if prev is None:
            continue
        prev_pPr = _ensure_pPr(prev)
        # Strip any pre-existing sectPr on this paragraph.
        for old in prev_pPr.findall(qn("w:sectPr")):
            prev_pPr.remove(old)
        prev_pPr.append(build_sectpr_for_digit(prev_digit))

    # Replace body-level sectPr with the LAST chapter's mapping.
    last_digit = chapters[-1][1]
    new_body_sect = build_sectpr_for_digit(last_digit)
    for old in body.findall(qn("w:sectPr")):
        body.remove(old)
    body.append(new_body_sect)

    # Remove marker paragraphs.
    for p in markers_to_remove:
        body.remove(p)

    # Remove RmdMarker* style definitions from styles.xml — they're only used
    # as transient digit-tagging anchors during post-processing.
    styles_el = doc.styles.element
    for s in list(styles_el.findall(qn("w:style"))):
        sid = s.get(qn("w:styleId")) or ""
        if re.match(r"^RmdMarker\d+$", sid):
            styles_el.remove(s)
            continue
        name_el = s.find(qn("w:name"))
        nval = name_el.get(qn("w:val")) if name_el is not None else ""
        if nval and re.match(r"^RmdMarker\d+$", nval):
            styles_el.remove(s)


def main() -> None:
    args = [a for a in sys.argv[1:] if a != "--quiet"]
    # --csl=<numeric|note|author-date>; default: numeric
    csl = "numeric"
    raw = []
    for a in args:
        if a.startswith("--csl="):
            csl = a.split("=", 1)[1]
        elif a in CSL_CHOICES:
            csl = a
        else:
            raw.append(a)
    if not raw:
        raw = ["docx"]
    pandoc = find_pandoc()
    formats = [fmt.replace("bookdown::", "").replace("_book", "") for fmt in raw]
    src = merge_sources(docx_mode=True)
    try:
        for f in formats:
            if f == "docx":
                render_docx(pandoc, src, lang="zh", csl=csl)
            elif f in ("docx-en", "docx_en"):
                render_docx(pandoc, src, lang="en", csl=csl)
            else:
                sys.exit(f"Unknown format: {f}")
    finally:
        if src.exists():
            try:
                src.unlink()
            except OSError:
                pass


if __name__ == "__main__":
    main()
